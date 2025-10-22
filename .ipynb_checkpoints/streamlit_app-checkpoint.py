# streamlit_app.py
"""
Art of Drawers Onboarding PDF Generator
Generates personalized onboarding PDFs for Designers and Installers
"""

import streamlit as st
import io
import base64
import re
from docx import Document
from cloudconvert import CloudConvert

# Import embedded templates
from templates_base64 import DESIGNER_DOCX_B64, INSTALLER_DOCX_B64


def normalize_name_for_greeting(first_name):
    """
    Convert first name to Title Case, preserving hyphenated capitalization.
    Example: "mary-jane" -> "Mary-Jane"
    """
    return first_name.title()


def normalize_name_for_email(name):
    """
    Strip non-alphanumeric characters and convert to lowercase.
    Example: "O'Brien-Smith" -> "obriensmith"
    """
    return re.sub(r'[^a-zA-Z0-9]', '', name).lower()


def generate_email(first_name, last_name):
    """
    Generate email: first_initial + last_name + @artofdrawers.com
    Example: Mat Fluker -> mfluker@artofdrawers.com
    """
    first_clean = normalize_name_for_email(first_name)
    last_clean = normalize_name_for_email(last_name)
    
    if not first_clean or not last_clean:
        raise ValueError("Names must contain at least one alphanumeric character")
    
    return f"{first_clean[0]}{last_clean}@artofdrawers.com"


def generate_canvas_username(first_name, last_name):
    """
    Generate Canvas username: firstname.lastname (lowercase, alphanumeric only)
    Example: Mat Fluker -> mat.fluker
    """
    first_clean = normalize_name_for_email(first_name)
    last_clean = normalize_name_for_email(last_name)
    
    if not first_clean or not last_clean:
        raise ValueError("Names must contain at least one alphanumeric character")
    
    return f"{first_clean}.{last_clean}"


def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace placeholders in a paragraph, handling Word's run splitting.
    """
    # Get the full text of the paragraph
    full_text = paragraph.text
    
    # Check if any placeholder exists in this paragraph
    needs_replacement = any(placeholder in full_text for placeholder in replacements.keys())
    
    if needs_replacement:
        # Perform all replacements
        new_text = full_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, value)
        
        # Clear all runs and add the new text as a single run
        # Preserve the style of the first run if it exists
        if paragraph.runs:
            first_run_font = paragraph.runs[0].font
            paragraph.clear()
            new_run = paragraph.add_run(new_text)
            # Copy font properties
            new_run.font.name = first_run_font.name
            new_run.font.size = first_run_font.size
            new_run.font.bold = first_run_font.bold
            new_run.font.italic = first_run_font.italic
        else:
            paragraph.clear()
            paragraph.add_run(new_text)


def replace_placeholders_in_docx(doc, replacements):
    """
    Replace all placeholders in the document (paragraphs, tables, headers, footers).
    """
    # Replace in main document paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)


def load_template(role):
    """
    Load the appropriate template from base64-encoded data.
    """
    if role == "Designer":
        template_b64 = DESIGNER_DOCX_B64
    else:  # Installer
        template_b64 = INSTALLER_DOCX_B64
    
    # Decode base64 to bytes
    template_bytes = base64.b64decode(template_b64)
    
    # Load into Document
    doc = Document(io.BytesIO(template_bytes))
    return doc


def generate_filled_docx(role, first_name, last_name):
    """
    Generate a filled .docx file in memory with placeholders replaced.
    """
    # Load template
    doc = load_template(role)
    
    # Generate values
    greeting = normalize_name_for_greeting(first_name) + ","
    email = generate_email(first_name, last_name)
    canvas_username = generate_canvas_username(first_name, last_name)
    
    # Define replacements
    replacements = {
        "{{GREETING}}": greeting,
        "{{GMAIL}}": email,
        "{{CANVAS_USERNAME}}": canvas_username
    }
    
    # Replace placeholders
    replace_placeholders_in_docx(doc, replacements)
    
    # Save to BytesIO
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()


def convert_docx_to_pdf(docx_bytes, api_key):
    """
    Convert .docx to .pdf using CloudConvert API.
    Returns PDF bytes.
    """
    cloudconvert = CloudConvert(api_key=api_key, sandbox=False)
    
    # Create a job
    job = cloudconvert.Job.create(payload={
        'tasks': {
            'upload-my-file': {
                'operation': 'import/upload'
            },
            'convert-my-file': {
                'operation': 'convert',
                'input': 'upload-my-file',
                'output_format': 'pdf',
                'input_format': 'docx'
            },
            'export-my-file': {
                'operation': 'export/url',
                'input': 'convert-my-file'
            }
        }
    })
    
    # Upload the file
    upload_task_id = job['tasks'][0]['id']
    upload_task = cloudconvert.Task.find(id=upload_task_id)
    cloudconvert.Task.upload(file_name='document.docx', task=upload_task, file=io.BytesIO(docx_bytes))
    
    # Wait for job completion
    job = cloudconvert.Job.wait(id=job['id'])
    
    # Download the result
    for task in job['tasks']:
        if task.get('operation') == 'export/url' and task.get('status') == 'finished':
            file_info = task['result']['files'][0]
            pdf_bytes = cloudconvert.download(filename=file_info['filename'], url=file_info['url'])
            return pdf_bytes
    
    raise Exception("PDF conversion failed")


def generate_pdf_filename(role, first_name, last_name):
    """
    Generate PDF filename: {role}-{first}_{last}-onboarding.pdf
    Example: designer-mat_fluker-onboarding.pdf
    """
    role_lower = role.lower()
    first_clean = normalize_name_for_email(first_name)
    last_clean = normalize_name_for_email(last_name)
    
    return f"{role_lower}-{first_clean}_{last_clean}-onboarding.pdf"


# Streamlit UI
def main():
    st.set_page_config(
        page_title="AOD Onboarding PDF Generator",
        page_icon="📄",
        layout="centered"
    )
    
    st.title("🎨 Art of Drawers")
    st.subheader("Onboarding PDF Generator")
    st.markdown("---")
    
    # Check for API key
    if "CLOUDCONVERT_API_KEY" not in st.secrets:
        st.error("⚠️ CloudConvert API key not found in secrets. Please configure it in Streamlit Cloud settings.")
        st.stop()
    
    # Form
    with st.form("onboarding_form"):
        st.markdown("### Employee Information")
        
        role = st.radio(
            "Role",
            options=["Designer", "Installer"],
            horizontal=True
        )
        
        col1, col2 = st.columns(2)
        with col1:
            first_name = st.text_input("First Name", placeholder="e.g., Mat")
        with col2:
            last_name = st.text_input("Last Name", placeholder="e.g., Fluker")
        
        submitted = st.form_submit_button("Generate Onboarding PDF", type="primary", use_container_width=True)
    
    # Process form submission
    if submitted:
        # Validation
        if not first_name.strip() or not last_name.strip():
            st.error("❌ Please enter both first and last name.")
            return
        
        try:
            with st.spinner("🔄 Generating your PDF..."):
                # Generate filled docx
                docx_bytes = generate_filled_docx(role, first_name.strip(), last_name.strip())
                
                # Convert to PDF
                api_key = st.secrets["CLOUDCONVERT_API_KEY"]
                pdf_bytes = convert_docx_to_pdf(docx_bytes, api_key)
                
                # Generate filename
                pdf_filename = generate_pdf_filename(role, first_name.strip(), last_name.strip())
                
                # Success message
                st.success("✅ PDF generated successfully!")
                
                # Preview info
                st.markdown("### 📋 Generated Details")
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Email:**")
                    st.code(generate_email(first_name.strip(), last_name.strip()))
                with col2:
                    st.write("**Canvas Username:**")
                    st.code(generate_canvas_username(first_name.strip(), last_name.strip()))
                
                # Download button
                st.download_button(
                    label="📥 Download Onboarding PDF",
                    data=pdf_bytes,
                    file_name=pdf_filename,
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True
                )
        
        except ValueError as e:
            st.error(f"❌ Invalid input: {str(e)}")
        except Exception as e:
            st.error(f"❌ Error generating PDF: {str(e)}")
            st.info("Please check your CloudConvert API key and quota, then try again.")


if __name__ == "__main__":
    main()